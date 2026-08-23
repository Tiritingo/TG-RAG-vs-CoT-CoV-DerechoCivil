#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extraer_texto_cc.py — Normaliza el corpus mixto de Corte Constitucional a texto plano.

Tras reparar_corpus_cc.py quedan 411 archivos en tres formatos (.rtf, .docx, .doc).
Este script los convierte a .txt uniforme para que el pipeline de indexacion
tenga una sola ruta de lectura, y verifica que cada salida tenga contenido real.

    pip install striprtf python-docx
    python extraer_texto_cc.py

Para los .doc (Word 97, formato OLE) hace falta un conversor externo. El script
busca, en orden: LibreOffice, Microsoft Word via COM, antiword. Si no encuentra
ninguno te lo dice y deja esos archivos pendientes; el resto se procesa igual.

Autor: Gerardo Aguilar — UPB, Maestría en Ciencia de Datos
"""
import csv, os, re, shutil, subprocess, sys, tempfile
from collections import Counter
from pathlib import Path

# ============================== CONFIG ==============================
RAIZ = Path(__file__).resolve().parent.parent.parent
ORIGEN = RAIZ / "01_Corpus_Raw" / "Sentencias" / "Corte_Constitucional" / "archivos_v3"
DESTINO = RAIZ / "01_Corpus_Raw" / "Sentencias" / "Corte_Constitucional" / "texto_plano"
MIN_CHARS = 2000            # una sentencia real siempre supera esto
# ====================================================================

RE_JUR = re.compile(
    r"\b(corte constitucional|sala plena|magistrad[oa] ponente|sentencia (?:C|T|SU)-|"
    r"acci[óo]n de tutela|constituci[óo]n pol[íi]tica|expediente|demandante|"
    r"RESUELVE|C[óo]digo Civil|contrato)\b", re.I)

RUTAS_SOFFICE = [
    "soffice",
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "/usr/bin/soffice", "/usr/bin/libreoffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
]


def hallar_soffice():
    for r in RUTAS_SOFFICE:
        if shutil.which(r) or Path(r).exists():
            return r
    return None


# ------------------------------------------------------------ extractores

def de_rtf(p):
    from striprtf.striprtf import rtf_to_text
    return rtf_to_text(p.read_text(encoding="utf-8", errors="ignore"), errors="ignore")


def de_docx(p):
    import docx
    d = docx.Document(str(p))
    partes = [x.text for x in d.paragraphs]
    for t in d.tables:                      # las sentencias traen tablas
        for fila in t.rows:
            partes.append("\t".join(c.text for c in fila.cells))
    return "\n".join(partes)


def de_doc_soffice(p, soffice):
    """Convierte .doc -> .docx con LibreOffice y luego extrae."""
    with tempfile.TemporaryDirectory() as tmp:
        r = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, str(p)],
            capture_output=True, timeout=180)
        salida = Path(tmp) / (p.stem + ".docx")
        if not salida.exists():
            raise RuntimeError("LibreOffice no produjo salida: %s"
                               % r.stderr.decode("utf-8", "replace")[:120])
        return de_docx(salida)


def de_doc_word(p):
    """Windows con Microsoft Word instalado."""
    import win32com.client  # pywin32
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(p.resolve()), ReadOnly=True)
        try:
            return doc.Content.Text
        finally:
            doc.Close(False)
    finally:
        word.Quit()


def de_doc_antiword(p):
    r = subprocess.run(["antiword", str(p)], capture_output=True, timeout=120)
    if r.returncode != 0:
        raise RuntimeError("antiword fallo")
    return r.stdout.decode("utf-8", "replace")


RE_HEXBLOB = re.compile(r"[0-9a-f]{200,}", re.I)


def limpiar(t):
    """
    Las imagenes incrustadas en un RTF van como bloques hexadecimales dentro
    de \\pict. striprtf no siempre los descarta y terminan en la salida: en
    C-583/15 eran 1,7 MB de basura sobre 1,8 MB de archivo. Se purgan aqui.
    """
    t = t.replace("\x00", " ").replace("\r", "\n")
    t = RE_HEXBLOB.sub(" ", t)
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


# ------------------------------------------------------------------ main

def main():
    if not ORIGEN.is_dir():
        sys.exit("No existe:\n  %s\nCorre primero reparar_corpus_cc.py" % ORIGEN)

    DESTINO.mkdir(parents=True, exist_ok=True)
    archivos = sorted(p for p in ORIGEN.iterdir()
                      if p.suffix.lower() in (".rtf", ".docx", ".doc"))
    if not archivos:
        sys.exit("No hay documentos en %s" % ORIGEN)

    soffice = hallar_soffice()
    tiene_doc = any(p.suffix.lower() == ".doc" for p in archivos)

    print("=" * 72)
    print("EXTRACCION DE TEXTO — %d documentos" % len(archivos))
    print("=" * 72)
    print("Origen : %s" % ORIGEN)
    print("Destino: %s" % DESTINO)
    if tiene_doc:
        print("Conversor .doc: %s" % (soffice or "NINGUNO — se intentara Word COM / antiword"))
    print()

    manifiesto = []
    stats = Counter()
    pendientes = []

    for n, p in enumerate(archivos, 1):
        ext = p.suffix.lower()
        salida = DESTINO / (p.stem + ".txt")

        if salida.exists() and salida.stat().st_size > MIN_CHARS:
            stats["ya_existia"] += 1
            continue

        texto, metodo, error = "", "", ""
        try:
            if ext == ".rtf":
                texto, metodo = de_rtf(p), "striprtf"
            elif ext == ".docx":
                texto, metodo = de_docx(p), "python-docx"
            else:
                if soffice:
                    texto, metodo = de_doc_soffice(p, soffice), "libreoffice"
                else:
                    for fn, nom in ((de_doc_word, "word-com"), (de_doc_antiword, "antiword")):
                        try:
                            texto, metodo = fn(p), nom
                            break
                        except Exception:
                            continue
                    if not texto:
                        raise RuntimeError("sin conversor para .doc")
        except Exception as e:
            error = str(e)[:100]

        texto = limpiar(texto or "")
        hits = len(RE_JUR.findall(texto))

        if error or len(texto) < MIN_CHARS or hits < 3:
            estado = "PENDIENTE" if error else "SOSPECHOSO"
            pendientes.append(p.name)
            stats[estado] += 1
            if not error:
                error = "%d chars, %d marcas juridicas" % (len(texto), hits)
        else:
            salida.write_text(texto, encoding="utf-8")
            estado = "OK"
            stats["OK"] += 1
            stats["fmt_" + ext] += 1

        manifiesto.append({"archivo": p.name, "formato": ext, "metodo": metodo,
                           "chars": len(texto), "marcas_juridicas": hits,
                           "estado": estado, "detalle": error})

        if n % 50 == 0:
            print("  %d/%d — %d ok, %d pendientes" % (n, len(archivos), stats["OK"], len(pendientes)))

    with open(DESTINO / "manifiesto_extraccion.csv", "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=["archivo", "formato", "metodo", "chars",
                                          "marcas_juridicas", "estado", "detalle"])
        w.writeheader()
        w.writerows(manifiesto)

    # ---------------------------------------------------------- resumen
    txts = sorted(DESTINO.glob("*.txt"))
    print("\n" + "=" * 72)
    print("RESULTADO")
    print("=" * 72)
    print("Extraidos OK   : %d" % stats["OK"])
    if stats["ya_existia"]:
        print("Ya existian    : %d" % stats["ya_existia"])
    print("Sospechosos    : %d" % stats["SOSPECHOSO"])
    print("Pendientes     : %d" % stats["PENDIENTE"])
    print("Archivos .txt  : %d" % len(txts))

    if txts:
        tam = sorted(t.stat().st_size for t in txts)
        print("Tamano texto   : min=%d  mediana=%d  max=%d chars"
              % (tam[0], tam[len(tam) // 2], tam[-1]))

    if pendientes:
        print("\nSin procesar (%d):" % len(pendientes))
        for x in pendientes[:12]:
            print("   %s" % x)
        if len(pendientes) > 12:
            print("   ... y %d mas" % (len(pendientes) - 12))
        if not soffice and tiene_doc:
            print("\nPara los .doc instala LibreOffice y vuelve a correr:")
            print("   https://es.libreoffice.org/descarga/")
            print("Se detecta solo; no hay que configurar nada.")

    if len(txts) == len(archivos):
        print("\nCorpus normalizado al 100%. Listo para reindexar.")
    print("\nManifiesto: %s" % (DESTINO / "manifiesto_extraccion.csv"))


if __name__ == "__main__":
    main()
